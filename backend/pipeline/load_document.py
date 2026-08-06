"""
Document ingestion (PDR Section 4.2 / Phase 4): turns a real compliance
document file (.txt, .pdf, or .docx) into plain text, ready to hand to
extract_claims(). Uses only free, open-source libraries.

Install dependencies (all free/open-source):
    pip install pdfplumber python-docx
"""
import os


def load_document_text(file_path: str) -> str:
    """Reads a compliance document from disk and returns its plain text,
    regardless of whether it's a .txt, .pdf, or .docx file.

    Raises ValueError for unsupported file types, and a clear error if a
    PDF appears to be a scanned image with no extractable text (OCR would
    be needed — a possible future extension, not handled here).
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return _load_txt(file_path)
    if ext == ".pdf":
        return _load_pdf(file_path)
    if ext == ".docx":
        return _load_docx(file_path)

    raise ValueError(
        f"Unsupported document type '{ext}'. Supported: .txt, .pdf, .docx"
    )


def _load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _load_pdf(file_path: str) -> str:
    import pdfplumber  # imported here so .txt-only usage doesn't require it installed

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise ValueError(
            f"No extractable text found in '{file_path}'. This usually means "
            f"the PDF is a scanned image rather than real text — OCR (e.g. "
            f"pytesseract) would be needed to read it, which isn't wired up yet."
        )

    return full_text


def _load_docx(file_path: str) -> str:
    import docx  # python-docx package, imported here for the same reason as above

    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of any tables, since compliance documents often
    # put key fields (policy number, claim amount, etc.) in table cells
    # rather than plain paragraphs.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)
